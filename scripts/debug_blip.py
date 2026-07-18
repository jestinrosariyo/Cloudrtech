"""
Standalone diagnostic: inspect <a:blip> tags inside a DOCX zip.
Run from the project root:
    .\venv\Scripts\python.exe scripts\debug_blip.py <path-to-docx>

If no path is given, it creates a minimal test DOCX with an embedded image
and inspects that.
"""
import sys, os, zipfile, re
from io import BytesIO

# Create a minimal test docx with an embedded image
def create_test_docx():
    from docx import Document
    from docx.shared import Pt
    from PIL import Image as PILImage

    # Create a tiny red 10x10 PNG
    img = PILImage.new('RGB', (10, 10), color='red')
    buf = BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)

    doc = Document()
    section = doc.sections[0]
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(buf, width=Pt(100), height=Pt(100))
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()


def inspect_docx(data, label='DOCX'):
    print(f'\n{"="*60}')
    print(f'  Inspecting: {label}')
    print(f'  Size: {len(data)} bytes')
    print(f'{"="*60}')

    with zipfile.ZipFile(BytesIO(data)) as z:
        for name in z.namelist():
            if name.endswith('.xml'):
                raw = z.read(name)
                text = raw.decode('utf-8', errors='replace')

                # Find all <a:blip ...> occurrences (self-closing or not)
                blip_tags = re.findall(r'<a:blip[^>]*/?>', text)
                if blip_tags:
                    print(f'\n  File: {name}')
                    for tag in blip_tags:
                        print(f'    {tag}')

                # Also check for blip without namespace prefix (fallback)
                blip_tags2 = re.findall(r'<[^>]*blip[^>]*/?>', text)
                if blip_tags2 and not blip_tags:
                    print(f'\n  File: {name} (non-prefixed)')
                    for tag in blip_tags2:
                        print(f'    {tag}')

    # Now test the current regex
    print(f'\n  --- Testing current regex pattern ---')
    pattern = r'<a:blip (r:embed="rId\d+")/>'
    replacement = r'<a:blip \1 cstate="print"/>'

    with zipfile.ZipFile(BytesIO(data)) as z:
        for name in z.namelist():
            if name.endswith('.xml'):
                raw = z.read(name)
                text = raw.decode('utf-8', errors='replace')
                matches = re.findall(pattern, text)
                if matches:
                    print(f'  File {name}: {len(matches)} match(es) with current regex')
                else:
                    # Check what's actually there
                    all_blips = re.findall(r'<a:blip[^>]*>', text)
                    if all_blips:
                        print(f'  File {name}: NO match with current regex')
                        print(f'  Actual <a:blip> tags found:')
                        for b in all_blips:
                            print(f'    {b}')

    # Test improved regex
    print(f'\n  --- Testing improved regex pattern ---')
    improved = r'(<a:blip\b[^>]*?)(\s*/>)'
    improved2 = r'(<a:blip\b[^>]*?)(\s*>)'

    with zipfile.ZipFile(BytesIO(data)) as z:
        for name in z.namelist():
            if name.endswith('.xml'):
                raw = z.read(name)
                text = raw.decode('utf-8', errors='replace')
                matches1 = re.findall(improved, text)
                matches2 = re.findall(improved2, text)
                if matches1 or matches2:
                    print(f'  File {name}: improved regex matched {len(matches1)} self-closing + {len(matches2)} non-self-closing')


if __name__ == '__main__':
    if len(sys.argv) > 1:
        path = sys.argv[1]
        if not os.path.exists(path):
            print(f'File not found: {path}')
            sys.exit(1)
        with open(path, 'rb') as f:
            data = f.read()
        inspect_docx(data, path)
    else:
        print('No file specified — creating a minimal test DOCX...')
        data = create_test_docx()
        inspect_docx(data, 'test_minimal.docx')
