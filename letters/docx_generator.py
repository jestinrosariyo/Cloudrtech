"""
Cloud R tech — DOCX Generator v5

Background method — PDF-style, repeats on EVERY page:
  - For all letter types (Appraisal, Offer, Contract, Experience):
    letterhead injected into the HEADER as a full-page anchored drawing
    (behindDoc=1, relativeFrom="page"). Since headers repeat on every page,
    the background appears on ALL pages including annexure — identical to
    how PDF's _draw_letterhead callback fires on every page.

  - For Payslip: NO background at all. Clean white document.
"""

import math, os, zipfile, re as _re
from io import BytesIO
from PIL import Image as PILImage
from PIL import ImageCms
from lxml import etree

from django.conf import settings

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Image paths ────────────────────────────────────────────────────────────────
def _img(fn):
    p = os.path.join(settings.BASE_DIR, 'static', 'letters', 'img', fn)
    return p if os.path.isfile(p) else None

def _letterhead_path():
    return _img('letterhead.png')

LETTERHEAD = _letterhead_path()
SIGNATURE  = _img('signature.png')
SEAL_STAMP = _img('seal_stamp.png')
PS_HEADER  = _img('payslip_header.jpg')

DEFAULT_HR = 'Raj Padmanaban'

# A4 page dimensions (pt) — matched to the PDF generator's measured template
PAGE_W  = 595.56
PAGE_H  = 842.52
LEFT_M  = 42.5
RIGHT_M = 79.6
TOP_M   = 152.5
BOT_M   = 99.0

# US Letter (payslip)
PS_PAGE_W  = 612.0
PS_PAGE_H  = 792.0
PS_LEFT_M  = 39.59
PS_RIGHT_M = 39.59
PS_TOP_M   = 36.0
PS_BOT_M   = 36.0


# ── PT → EMU (1 pt = 12700 EMU) ───────────────────────────────────────────────
def _emu(pt):
    return int(pt * 12700)


# ── ICC-aware JPEG helper ─────────────────────────────────────────────────────
# Embeds an sRGB ICC profile so Word displays colours exactly as the PDF does.
# Every output JPEG is tagged sRGB — this is what fixes the colour shift.
def _to_icc_free_jpeg(path_or_buf, quality=98):
    img = PILImage.open(path_or_buf).convert('RGB')
    srgb = ImageCms.ImageCmsProfile(ImageCms.createProfile('sRGB'))
    buf = BytesIO()
    img.save(buf, format='JPEG', quality=quality, subsampling=0,
             icc_profile=srgb.tobytes())
    buf.seek(0)
    return buf, img.size


# ── Post-process: add cstate='print' to every <a:blip> ────────────────────────
def _postprocess_docx_color(docx_bytes):
    buf_in  = BytesIO(docx_bytes)
    buf_out = BytesIO()
    with zipfile.ZipFile(buf_in, 'r') as zin, \
         zipfile.ZipFile(buf_out, 'w', zipfile.ZIP_DEFLATED) as zout:
        for name in zin.namelist():
            data = zin.read(name)
            if name.endswith('.xml'):
                content = data.decode('utf-8')
                content = _re.sub(
                    r'<a:blip (r:embed="rId\d+")(\/?)>',
                    r'<a:blip \1 cstate="print"\2>',
                    content
                )
                data = content.encode('utf-8')
            zout.writestr(name, data)
    buf_out.seek(0)
    return buf_out.getvalue()


# ── Full-page background — DrawingML anchor in BODY (behind text) ──
# Word intentionally fades header/footer images in Print Layout view,
# which cannot be overridden via XML attributes.  To render at FULL
# opacity the image MUST live in the document body, anchored at
# page-relative position (0,0) with behindDoc=1.
# A separate call is needed after each add_page_break() so every page
# gets its own background paragraph.
_BG_NSMAP = {
    'w':   'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'wp':  'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a':   'http://schemas.openxmlformats.org/drawingml/2006/main',
    'r':   'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
}
_BG_W, _BG_WP, _BG_A, _BG_R, _BG_PIC = [_BG_NSMAP[k] for k in ('w', 'wp', 'a', 'r', 'pic')]


def _make_bg_paragraph(rId, page_w_pt, page_h_pt, doc_pr_id=100):
    """Return a zero-height <w:p> containing a full-page behind-text anchor."""
    drawing = etree.Element(f'{{{_BG_W}}}drawing', nsmap=_BG_NSMAP)

    anchor = etree.SubElement(drawing, f'{{{_BG_WP}}}anchor')
    for attr, val in [('distT','0'),('distB','0'),('distL','0'),('distR','0'),
                      ('simplePos','0'),('relativeHeight','0'),('behindDoc','1'),
                      ('locked','0'),('layoutInCell','1'),('allowOverlap','1')]:
        anchor.set(attr, val)

    sp = etree.SubElement(anchor, f'{{{_BG_WP}}}simplePos')
    sp.set('x', '0'); sp.set('y', '0')

    pH = etree.SubElement(anchor, f'{{{_BG_WP}}}positionH')
    pH.set('relativeFrom', 'page')
    offH = etree.SubElement(pH, f'{{{_BG_WP}}}posOffset')
    offH.text = '0'

    pV = etree.SubElement(anchor, f'{{{_BG_WP}}}positionV')
    pV.set('relativeFrom', 'page')
    offV = etree.SubElement(pV, f'{{{_BG_WP}}}posOffset')
    offV.text = '0'

    extent = etree.SubElement(anchor, f'{{{_BG_WP}}}extent')
    extent.set('cx', str(_emu(page_w_pt)))
    extent.set('cy', str(_emu(page_h_pt)))

    etree.SubElement(anchor, f'{{{_BG_WP}}}wrapNone')

    docPr = etree.SubElement(anchor, f'{{{_BG_WP}}}docPr')
    docPr.set('id', str(doc_pr_id))
    docPr.set('name', 'Background')

    etree.SubElement(anchor, f'{{{_BG_WP}}}cNvGraphicFramePr')

    graphic = etree.SubElement(anchor, f'{{{_BG_A}}}graphic')
    gData = etree.SubElement(graphic, f'{{{_BG_A}}}graphicData')
    gData.set('uri', _BG_PIC)

    pic = etree.SubElement(gData, f'{{{_BG_PIC}}}pic')
    nvPicPr = etree.SubElement(pic, f'{{{_BG_PIC}}}nvPicPr')
    cNvPr = etree.SubElement(nvPicPr, f'{{{_BG_PIC}}}cNvPr')
    cNvPr.set('id', str(doc_pr_id))
    cNvPr.set('name', 'Background')
    etree.SubElement(nvPicPr, f'{{{_BG_PIC}}}cNvPicPr')

    blipFill = etree.SubElement(pic, f'{{{_BG_PIC}}}blipFill')
    blip = etree.SubElement(blipFill, f'{{{_BG_A}}}blip')
    blip.set(f'{{{_BG_R}}}embed', rId)
    stretch = etree.SubElement(blipFill, f'{{{_BG_A}}}stretch')
    etree.SubElement(stretch, f'{{{_BG_A}}}fillRect')

    spPr = etree.SubElement(pic, f'{{{_BG_PIC}}}spPr')
    xfrm = etree.SubElement(spPr, f'{{{_BG_A}}}xfrm')
    o = etree.SubElement(xfrm, f'{{{_BG_A}}}off')
    o.set('x', '0'); o.set('y', '0')
    e = etree.SubElement(xfrm, f'{{{_BG_A}}}ext')
    e.set('cx', str(_emu(page_w_pt)))
    e.set('cy', str(_emu(page_h_pt)))
    pg = etree.SubElement(spPr, f'{{{_BG_A}}}prstGeom')
    pg.set('prst', 'rect')

    run = OxmlElement('w:r')
    run.append(drawing)

    bg_para = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    sp_el = OxmlElement('w:spacing')
    sp_el.set(qn('w:before'), '0')
    sp_el.set(qn('w:after'),  '0')
    sp_el.set(qn('w:line'),   '0')
    sp_el.set(qn('w:lineRule'), 'exact')
    pPr.append(sp_el)
    bg_para.append(pPr)
    bg_para.append(run)
    return bg_para


def _add_body_background(doc, image_path, page_w_pt, page_h_pt, doc_pr_id=100):
    """Insert a full-page background image into the body at the current position."""
    if not image_path:
        return
    img_buf, _ = _to_icc_free_jpeg(image_path)
    rId, _ = doc.part.get_or_add_image(img_buf)
    bg_para = _make_bg_paragraph(rId, page_w_pt, page_h_pt, doc_pr_id)
    body = doc.element.body
    body.insert(0, bg_para)


def _add_body_background_after_break(doc, image_path, page_w_pt, page_h_pt, doc_pr_id=200):
    """Insert a full-page background image at the current end of the document
    content (typically after a page break, or mid-flow as a safety checkpoint).

    IMPORTANT: w:sectPr must always remain the LAST child of w:body — appending
    after it is invalid OOXML and causes unpredictable pagination/blank pages
    in Word/LibreOffice. We insert immediately before it instead.
    """
    if not image_path:
        return
    img_buf, _ = _to_icc_free_jpeg(image_path)
    rId, _ = doc.part.get_or_add_image(img_buf)
    bg_para = _make_bg_paragraph(rId, page_w_pt, page_h_pt, doc_pr_id)
    body = doc.element.body
    sectPr = body.find(qn('w:sectPr'))
    if sectPr is not None:
        sectPr.addprevious(bg_para)
    else:
        body.append(bg_para)


def _add_anchor_image_to_paragraph(paragraph, rId, width_pt, height_pt):
    """Add an inline DrawingML image to a paragraph.
    Uses <wp:inline> so the image is part of the document flow,
    expanding the paragraph/cell to contain it with no white gaps."""
    NSMAP = {
        'w':   'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'wp':  'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a':   'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r':   'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    }
    W, WP, A, R, PIC = [NSMAP[k] for k in ('w', 'wp', 'a', 'r', 'pic')]

    drawing = etree.Element(f'{{{W}}}drawing', nsmap=NSMAP)

    inline = etree.SubElement(drawing, f'{{{WP}}}inline')
    for attr, val in [('distT','0'),('distB','0'),('distL','0'),('distR','0')]:
        inline.set(attr, val)

    extent = etree.SubElement(inline, f'{{{WP}}}extent')
    extent.set('cx', str(_emu(width_pt)))
    extent.set('cy', str(_emu(height_pt)))

    docPr = etree.SubElement(inline, f'{{{WP}}}docPr')
    docPr.set('id', '50')
    docPr.set('name', 'PayslipBanner')

    etree.SubElement(inline, f'{{{WP}}}cNvGraphicFramePr')

    graphic = etree.SubElement(inline, f'{{{A}}}graphic')
    gData = etree.SubElement(graphic, f'{{{A}}}graphicData')
    gData.set('uri', PIC)

    pic = etree.SubElement(gData, f'{{{PIC}}}pic')
    nvPicPr = etree.SubElement(pic, f'{{{PIC}}}nvPicPr')
    cNvPr = etree.SubElement(nvPicPr, f'{{{PIC}}}cNvPr')
    cNvPr.set('id', '50')
    cNvPr.set('name', 'PayslipBanner')
    etree.SubElement(nvPicPr, f'{{{PIC}}}cNvPicPr')

    blipFill = etree.SubElement(pic, f'{{{PIC}}}blipFill')
    blip = etree.SubElement(blipFill, f'{{{A}}}blip')
    blip.set(f'{{{R}}}embed', rId)
    stretch = etree.SubElement(blipFill, f'{{{A}}}stretch')
    etree.SubElement(stretch, f'{{{A}}}fillRect')

    spPr = etree.SubElement(pic, f'{{{PIC}}}spPr')
    xfrm = etree.SubElement(spPr, f'{{{A}}}xfrm')
    o = etree.SubElement(xfrm, f'{{{A}}}off')
    o.set('x', '0'); o.set('y', '0')
    e = etree.SubElement(xfrm, f'{{{A}}}ext')
    e.set('cx', str(_emu(width_pt)))
    e.set('cy', str(_emu(height_pt)))
    pg = etree.SubElement(spPr, f'{{{A}}}prstGeom')
    pg.set('prst', 'rect')

    run = OxmlElement('w:r')
    run.append(drawing)
    paragraph._p.append(run)


# ── Document setup ─────────────────────────────────────────────────────────────
def _setup_letter_doc():
    """
    A4 doc. Background is NOT added here — it is added after content is
    built, via _finalise_letter(), so it goes into the header correctly.
    """
    doc = Document()
    section = doc.sections[0]
    section.page_width    = Pt(PAGE_W)
    section.page_height   = Pt(PAGE_H)
    section.left_margin   = Pt(LEFT_M)
    section.right_margin  = Pt(RIGHT_M)
    section.top_margin    = Pt(TOP_M)
    section.bottom_margin = Pt(BOT_M)
    section.header_distance = Pt(0)
    section.footer_distance = Pt(0)

    # Start with empty header/footer (background added later)
    header = section.header
    header.is_linked_to_previous = False
    if header.paragraphs:
        hp = header.paragraphs[0]; hp.clear()
        _para_spacing(hp, before=0, after=0)

    footer = section.footer
    footer.is_linked_to_previous = False
    if footer.paragraphs:
        fp = footer.paragraphs[0]; fp.clear()
        _para_spacing(fp, before=0, after=0)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.widow_control = True

    # python-docx's built-in template ships with compatibilityMode=14 (Word 2010),
    # which forces every generated file to open in "Compatibility Mode" in Word.
    # Bump to 15 (Word 2013+ native) so the file opens normally.
    compat = doc.settings.element.find(qn('w:compat'))
    if compat is not None:
        for cs in compat.findall(qn('w:compatSetting')):
            if cs.get(qn('w:name')) == 'compatibilityMode':
                cs.set(qn('w:val'), '15')

    return doc


def _finalise_letter(doc, page_w=PAGE_W, page_h=PAGE_H, bg_path=None):
    """
    Serialise the document. Background is now handled by body-anchored
    images (see _add_body_background / _add_body_background_after_break)
    instead of the header, to avoid Word's Print Layout view opacity fade.
    """
    buf = BytesIO()
    doc.save(buf)
    return _postprocess_docx_color(buf.getvalue())


# ── Low-level helpers ──────────────────────────────────────────────────────────
def _salutation(gender):
    return 'Mr.' if str(gender).strip().lower() in ('male','mr','mr.','m') else 'Ms.'

def _para_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:spacing')): pPr.remove(old)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(int(before * 20)))
    sp.set(qn('w:after'),  str(int(after  * 20)))
    if line is not None:
        sp.set(qn('w:line'),     str(int(line * 20)))
        sp.set(qn('w:lineRule'), 'exact')
    pPr.append(sp)

def _run(para, text, bold=False, italic=False, size=12):
    r = para.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size)
    r.font.name = 'Times New Roman'
    r.font.color.rgb = RGBColor(0, 0, 0)
    rPr = r._r.get_or_add_rPr()
    rf  = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rPr.insert(0, rf)
    rf.set(qn('w:ascii'), 'Times New Roman')
    rf.set(qn('w:hAnsi'), 'Times New Roman')
    return r

def _add_para(doc, text='', align=WD_ALIGN_PARAGRAPH.LEFT,
              bold=False, italic=False, size=12,
              before=0, after=0, line=17.5):
    p = doc.add_paragraph()
    p.alignment = align
    _para_spacing(p, before=before, after=after, line=line)
    if text: _run(p, text, bold=bold, italic=italic, size=size)
    return p

def _cell_borders(cell, top=True, bottom=True, left=True, right=True):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')): tcPr.remove(old)
    tcB = OxmlElement('w:tcBorders')
    for side, show in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        el = OxmlElement(f'w:{side}')
        if show:
            el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'4')
            el.set(qn('w:space'),'0');   el.set(qn('w:color'),'000000')
        else:
            el.set(qn('w:val'),'none');  el.set(qn('w:sz'),'0')
            el.set(qn('w:color'),'auto')
        tcB.append(el)
    tcPr.append(tcB)

def _cell_margins(cell, top=0, bottom=0, left=60, right=60):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcMar')): tcPr.remove(old)
    tcM = OxmlElement('w:tcMar')
    for side, v in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(v)); el.set(qn('w:type'),'dxa')
        tcM.append(el)
    tcPr.append(tcM)

def _col_width(cell, pt):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcW')): tcPr.remove(old)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(pt*20))); tcW.set(qn('w:type'),'dxa')
    tcPr.append(tcW)

def _fixed_table_layout(table, total_width_pt, col_widths_pt=None):
    """
    Force Word/LibreOffice to honour explicit cell widths instead of
    autofit-to-content (which causes wrapping like 'Management Allowance'
    breaking onto 2 lines even though the column is wide enough).

    Critical: tblLayout=fixed alone is NOT enough — Word/LibreOffice also
    require a <w:tblGrid> with matching <w:gridCol> widths, otherwise they
    fall back to content-based autofit regardless of tblLayout.
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    for old in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(old)
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)

    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(total_width_pt * 20)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    # Disable autofit explicitly
    for old in tblPr.findall(qn('w:tblOverlap')):
        tblPr.remove(old)

    # Build/replace tblGrid
    if col_widths_pt:
        existing_grid = tbl.find(qn('w:tblGrid'))
        if existing_grid is not None:
            tbl.remove(existing_grid)
        grid = OxmlElement('w:tblGrid')
        for w in col_widths_pt:
            gc = OxmlElement('w:gridCol')
            gc.set(qn('w:w'), str(int(w * 20)))
            grid.append(gc)
        # tblGrid must come right after tblPr
        tblPr.addnext(grid)

def _neg_indent(para, left_pt):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:ind')): pPr.remove(old)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(int(-left_pt * 20)))
    pPr.append(ind)


# ── Salary calculations ────────────────────────────────────────────────────────
def _salary_calcs(m):
    m = float(m)
    basic = round(m*0.40,2); hra = round(m*0.20,2); med = round(m*0.10,2)
    perf  = round(m*0.15,2); conv= round(m*0.10,2); mgmt= round(m*0.05,2)
    ctc   = round(basic+hra+med+perf+conv+mgmt,2)
    ded   = math.ceil(ctc*0.05/100)*100
    net   = round(ctc-ded,2)
    return basic,hra,med,perf,conv,mgmt,ctc,ded,net

def _fm(v): return f"{v:,.2f}"
def _fa(v):
    s=f"{v:.2f}"; i,d=s.split('.')
    if len(i)<=3: return f"{i}.{d}"
    r=i[-3:]; i=i[:-3]
    while i: r=i[-2:]+','+r; i=i[:-2]
    return r.lstrip(',')+'.'+d


# ── Annexure ───────────────────────────────────────────────────────────────────
def _annexure(doc, sal, name, designation, ctc_lpa, monthly_ctc):
    p = doc.add_paragraph(); _para_spacing(p, before=0, after=8, line=17.5)
    _run(p,'Annexure - (Salary Details)',bold=True,italic=True)

    info=[('Name',f'{sal} {name}'),('Designation',designation),('Cost to the company',f'{ctc_lpa} LPA')]
    t=doc.add_table(rows=len(info),cols=3); t.alignment=WD_TABLE_ALIGNMENT.LEFT
    cw=[141,36,200]
    for i,(lbl,val) in enumerate(info):
        for j,cell in enumerate(t.rows[i].cells):
            _col_width(cell,cw[j]); _cell_borders(cell,False,False,False,False)
            _cell_margins(cell,0,0,0,0)
            p2=cell.paragraphs[0]; p2.clear(); _para_spacing(p2,before=0,after=5,line=17.5)
            _run(p2, lbl if j==0 else (':' if j==1 else val))

    sp=doc.add_paragraph(); _para_spacing(sp,before=0,after=28)
    _add_body_background_after_break(doc, LETTERHEAD, PAGE_W, PAGE_H, doc_pr_id=151)

    # Safety checkpoint: if the salary table below pushes onto a further
    # page, that overflow page still needs its own background.
    _add_body_background_after_break(doc, LETTERHEAD, PAGE_W, PAGE_H, doc_pr_id=151)

    basic,hra,med,perf,conv,mgmt,ctc,ded,net=_salary_calcs(monthly_ctc)
    rows=[
        ('Component',           None,    None,     True, True),
        ('Basic Salary',        basic,   basic*12, False,False),
        ('House Rent Allowance (HRA)', hra, hra*12,False,False),
        ('Medical Allowance',   med,     med*12,   False,False),
        ('Performance Bonus',   perf,    perf*12,  False,False),
        ('Conveyance',          conv,    conv*12,  False,False),
        ('Allowance Management',mgmt,    mgmt*12,  False,False),
        ('Total Compensation (CTC)',ctc, ctc*12,   True, True),
        ('Deduction',           ded,     ded*12,   False,False),
        ('Net Salary',          net,     net*12,   True, True),
    ]
    st=doc.add_table(rows=len(rows),cols=3); st.alignment=WD_TABLE_ALIGNMENT.LEFT
    sw=[165,60,85]
    L=WD_ALIGN_PARAGRAPH.LEFT; R=WD_ALIGN_PARAGRAPH.RIGHT; C=WD_ALIGN_PARAGRAPH.CENTER
    for i,(lbl,mv,av,bold,italic) in enumerate(rows):
        row=st.rows[i]
        txts=[lbl, 'Monthly' if i==0 else _fm(mv), 'Annual' if i==0 else _fa(av)]
        alns=[L, C if i==0 else R, C if i==0 else R]
        for j,(cell,txt,aln) in enumerate(zip(row.cells,txts,alns)):
            _col_width(cell,sw[j]); _cell_borders(cell)
            _cell_margins(cell,2,2,60 if j==0 else 20,40)
            p3=cell.paragraphs[0]; p3.clear(); p3.alignment=aln
            _para_spacing(p3,before=2,after=2,line=15)
            _run(p3,txt,bold=bold,italic=italic)


# ── Floating (anchored) picture helper ──────────────────────────────────────
# python-docx only inserts INLINE pictures via add_picture(). The PDF's
# signature+seal are absolutely positioned and overlap (seal stamped over
# the bottom-right corner of the signature). To match that in Word we
# convert the inline drawing into a floating <wp:anchor> positioned
# relative to the paragraph, in EMU, exactly like the PDF's pt offsets.
def _float_picture(paragraph, image_path, x_pt, y_pt, width_pt, height_pt,
                    behind=False, z=0):
    WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    run = paragraph.add_run()
    if isinstance(image_path, BytesIO):
        image_path.seek(0)
    run.add_picture(image_path, width=Pt(width_pt), height=Pt(height_pt))

    drawing = run._r.find(qn('w:drawing'))
    inline  = drawing.find(qn('wp:inline'))
    extent  = inline.find(qn('wp:extent'))
    docPr   = inline.find(qn('wp:docPr'))
    graphic = inline.find(qn('a:graphic'))
    inline.remove(extent); inline.remove(docPr); inline.remove(graphic)

    anchor = OxmlElement('wp:anchor')
    anchor.set('behindDoc', '1' if behind else '0')
    anchor.set('distT', '0'); anchor.set('distB', '0')
    anchor.set('distL', '0'); anchor.set('distR', '0')
    anchor.set('simplePos', '0')
    anchor.set('locked', '0')
    anchor.set('layoutInCell', '1')
    anchor.set('allowOverlap', '1')
    anchor.set('relativeHeight', str(251650000 + z))

    simplePos = OxmlElement('wp:simplePos')
    simplePos.set('x', '0'); simplePos.set('y', '0')
    anchor.append(simplePos)

    posH = OxmlElement('wp:positionH')
    posH.set('relativeFrom', 'column')  # horizontal equiv. of "paragraph" — ST_RelFromH has no "paragraph" value
    offH = OxmlElement('wp:posOffset'); offH.text = str(_emu(x_pt))
    posH.append(offH); anchor.append(posH)

    posV = OxmlElement('wp:positionV')
    posV.set('relativeFrom', 'paragraph')
    offV = OxmlElement('wp:posOffset'); offV.text = str(_emu(y_pt))
    posV.append(offV); anchor.append(posV)

    anchor.append(extent)
    anchor.append(OxmlElement('wp:wrapNone'))
    anchor.append(docPr)
    anchor.append(OxmlElement('wp:cNvGraphicFramePr'))
    anchor.append(graphic)

    drawing.remove(inline)
    drawing.append(anchor)
    return run


# ── Signature block ────────────────────────────────────────────────────────────
def _sig_block(doc, hr_name=None, sig_src=None):
    name = hr_name or DEFAULT_HR
    pf=doc.add_paragraph(); _para_spacing(pf,before=0,after=6,line=17.5)
    _run(pf,'For '); _run(pf,'M/s Cloud R tech',bold=True,italic=True)

    # Anchor host paragraph — matches the PDF flowable's 200x104.2pt box.
    # Reserve vertical space with an empty run so subsequent paragraphs
    # (name, HR Manager) don't collide with the floating images.
    host = doc.add_paragraph()
    _para_spacing(host, before=0, after=0, line=None)
    host_run = host.add_run()
    host_run.font.size = Pt(1)
    pPr = host._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing'); pPr.append(spacing)
    spacing.set(qn('w:after'), str(int(104.2 * 20)))  # reserve ~104.2pt

    sig_path = sig_src if sig_src is not None else SIGNATURE
    if sig_path:
        try:
            # PDF: signature x=2.9pt, y=40.0pt from top of the 104.2pt box, 80.2x51.8pt
            _float_picture(host, sig_path, x_pt=2.9, y_pt=40.0,
                            width_pt=80.2, height_pt=51.8, z=0)
        except Exception:
            pass

    if SEAL_STAMP:
        try:
            # PDF: seal x=60.7pt, y=11.8pt from top of the box, 82.8x83.4pt — overlaps signature
            _float_picture(host, SEAL_STAMP, x_pt=60.7, y_pt=11.8,
                            width_pt=82.8, height_pt=83.4, z=1)
        except Exception:
            pass

    pn=doc.add_paragraph(); _para_spacing(pn,before=4,after=0,line=17.5)
    _run(pn,name,bold=True,italic=True)
    pr=doc.add_paragraph(); _para_spacing(pr,before=0,after=0,line=17.5)
    _run(pr,'HR Manager',bold=True,italic=True)


# ═══════════════════════════════════════════════════════════════════════════════
# 1. APPRAISAL LETTER
# ═══════════════════════════════════════════════════════════════════════════════
def generate_appraisal(data):
    doc=_setup_letter_doc()
    _add_body_background(doc, LETTERHEAD, PAGE_W, PAGE_H, doc_pr_id=100)
    name=data['name']; sal=_salutation(data.get('gender','male'))
    address=data['address']; pincode=data['pincode']; date_str=data['date']
    current=float(data['current_monthly']); new_m=float(data['new_monthly'])
    designation=data['designation']
    ctc_lpa=data.get('ctc_lpa','').replace('LPA','').replace('lpa','').strip()
    hr_name=data.get('hr_name'); sig_src=data.get('sig_src')

    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p,before=0,after=20,line=17.5); _run(p,'APPRAISAL LETTER',bold=True,italic=True)

    p2=doc.add_paragraph(); _para_spacing(p2,before=0,after=10,line=15.7); _run(p2,'To',bold=True,size=11.52)
    p3=doc.add_paragraph(); _para_spacing(p3,before=0,after=0,line=15.7); _run(p3,f'{sal} {name},',size=11.52)
    p4=doc.add_paragraph(); _para_spacing(p4,before=0,after=0,line=15.7); _run(p4,address,size=11.52)
    p5=doc.add_paragraph(); _para_spacing(p5,before=0,after=30,line=15.7); _run(p5,f'Pincode - {pincode}',size=11.52)
    p6=doc.add_paragraph(); _para_spacing(p6,before=0,after=14,line=17.5); _run(p6,f'Dear {sal} {name},')

    p7=doc.add_paragraph(); p7.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    _para_spacing(p7,before=0,after=14,line=17.5)
    _run(p7,'This Appraisal Letter is made on ')
    _run(p7,date_str,bold=True,italic=True)
    _run(p7,f', to appreciate your performance. We revise the remuneration from INR '
         f'{current:,.0f}/- per month to INR {new_m:,.0f}/- per month with effect from beginning '
         f'of the next month. A detailed salary structure is given in the annexure.')

    p8=doc.add_paragraph(); _para_spacing(p8,before=0,after=40,line=17.5)
    _run(p8,'All the other terms and conditions of the original contract remain unchanged.')

    _sig_block(doc,hr_name=hr_name,sig_src=sig_src)

    # Use a run-level page break (inside the last paragraph) instead of
    # doc.add_page_break() which creates an extra empty paragraph that
    # can push content onto a third page.
    last_body_para = doc.paragraphs[-1]
    run_br = last_body_para.add_run()
    run_br.add_break(WD_BREAK.PAGE)

    _add_body_background_after_break(doc, LETTERHEAD, PAGE_W, PAGE_H, doc_pr_id=101)
    _annexure(doc,sal,name,designation,ctc_lpa,new_m)
    # Background on every page via body-anchored images
    return _finalise_letter(doc, bg_path=LETTERHEAD)


# ═══════════════════════════════════════════════════════════════════════════════
# 2. EXPERIENCE / RELIEVING LETTER
# ═══════════════════════════════════════════════════════════════════════════════
def generate_experience(data):
    doc=_setup_letter_doc()
    _add_body_background(doc, LETTERHEAD, PAGE_W, PAGE_H, doc_pr_id=100)
    name=data['name']; sal=_salutation(data.get('gender','male'))
    ref=data['ref']; closing=data['closing_date']; designation=data['designation']
    doj=data['date_of_joining']; dor=data['date_of_relieving']
    reason=data['reason']; conduct=data['conduct']
    hr_name=data.get('hr_name'); sig_src=data.get('sig_src')

    p=doc.add_paragraph(); _para_spacing(p,before=0,after=14,line=17.5)
    _run(p,'Ref : '); _run(p,ref,bold=True,italic=True)

    p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p2,before=0,after=14,line=17.5)
    _run(p2,'Sub: Relieving / Experience Letter',bold=True,italic=True)

    p3=doc.add_paragraph(); _para_spacing(p3,before=0,after=14,line=17.5)
    _run(p3,f'Dear {sal} {name},',bold=True)

    p4=doc.add_paragraph(); p4.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    _para_spacing(p4,before=0,after=18,line=17.5)
    _run(p4,'This is with reference to your resignation from the services of this Organization. '
         'Your resignation has been accepted and you are relieved from the services of the '
         'Organization at the closing of business hours on ')
    _run(p4,closing,bold=True,italic=True); _run(p4,'.')

    det=[('Designation',designation,False),('Date of Joining',doj,False),
         ('Date of Relieving',dor,False),('Reason for Relieving',reason,False),('Conduct',conduct,True)]
    t=doc.add_table(rows=len(det),cols=3); t.alignment=WD_TABLE_ALIGNMENT.LEFT
    cw=[141,36,200]
    for i,(lbl,val,vi) in enumerate(det):
        for j,cell in enumerate(t.rows[i].cells):
            _col_width(cell,cw[j]); _cell_borders(cell,False,False,False,False); _cell_margins(cell,0,0,0,0)
            p5=cell.paragraphs[0]; p5.clear(); _para_spacing(p5,before=0,after=5,line=17.5)
            _run(p5, lbl if j==0 else (':' if j==1 else val), italic=(vi if j==2 else False))

    sp=doc.add_paragraph(); _para_spacing(sp,before=18,after=14)
    p6=doc.add_paragraph(); _para_spacing(p6,before=0,after=14,line=17.5)
    _run(p6,'We wish you all the best in your future endeavors.')
    _sig_block(doc,hr_name=hr_name,sig_src=sig_src)
    return _finalise_letter(doc, bg_path=LETTERHEAD)


# ═══════════════════════════════════════════════════════════════════════════════
# 3. OFFER LETTER
# ═══════════════════════════════════════════════════════════════════════════════
def generate_offer(data):
    doc=_setup_letter_doc()
    _add_body_background(doc, LETTERHEAD, PAGE_W, PAGE_H, doc_pr_id=100)
    date_str=data['date']; name=data['name']; sal=_salutation(data.get('gender','male'))
    address=data['address']; pincode=data['pincode']; designation=data['designation']
    joining=data['joining_date']; ctc_monthly=data['ctc_monthly']
    ctc_lpa=data['ctc_lpa'].replace('LPA','').replace('lpa','').strip()
    hr_name=data.get('hr_name'); sig_src=data.get('sig_src')

    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    _para_spacing(p,before=0,after=10,line=17.5); _run(p,f'Date : {date_str}',bold=True)

    _add_para(doc,'To',bold=True,before=0,after=5)
    _add_para(doc,f'{sal} {name} ,',before=0,after=0)
    _add_para(doc,address,before=0,after=0)
    _add_para(doc,f'Pincode - {pincode}',before=0,after=14)

    p2=doc.add_paragraph(); _para_spacing(p2,before=0,after=10,line=17.5)
    _run(p2,'Dear '); _run(p2,f'{sal} {name} ,',bold=True)

    p3=doc.add_paragraph(); p3.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    _para_spacing(p3,before=0,after=10,line=17.5)
    _run(p3,'We are pleased to appoint you to the position of ')
    _run(p3,designation,bold=True); _run(p3,' with ')
    _run(p3,'M/s Cloud R tech',bold=True)
    _run(p3,' and you will be working with us with the following terms and conditions')

    clauses=[
        [('Your Joining period will commence on ',False,False),(joining,True,False),(' on which you will be reporting to HR person at 10am in the previously communicated address',False,False)],
        [(f'Your annual total compensation will be INR ',False,False),(ctc_lpa,True,False),(' LPA. Apart from this, you are eligible for shift allowance depending on the project cost. Please find the salary structure given in the annexure.',False,False)],
        [("Your present place of work will be at Chennai, but during the course of the above assignment, you shall be liable to be posted / transferred anywhere to serve any of the Company's Projects or any other establishment in India or outside, at the sole discretion of the Management.",False,False)],
        [('You will be eligible for 12 working days of vacation and public holidays as notified by the company.',False,False)],
        [('This appointment is subject to you, having been found medically (physically and mentally) fit by the authorized Medical Practitioner.',False,False)],
        [("You will not (except in the normal course of the Company's business) publish any article or statement, deliver any lecture or broadcast or make any communication to the press, including magazine publication relating to the Company's products or to any matter with which the Company may be concerned, unless you have previously applied to and obtained.",False,False)],
        [('Any of our technical or other important information which might come into your possession during the continuance of your assignment with us shall not be disclosed, divulged or made public by you even thereafter.',False,False)],
        [("If at any time in our opinion, which is final in this matter you are found non-performer or guilty of fraud, dishonest, disobedience, disorderly behavior, negligence, indiscipline, absence from duty without permission or any other conduct considered by us deterrent to our interest or of violation of one or more terms of this letter, your services may be terminated without notice and on account of reason of any of the acts or omission, the company shall be entitled to recover the damages from you.",False,False)],
        [('All appointments are based on the information furnished by you in your employment application and all further declarations and undertakings. Hence, any false statement or information furnished as above will lead to your dismissal without notice.',False,False)],
        [('You hereby warrant that you are not in breach of any contract with any third party or restricted in any way in your ability to undertake or perform the duties of your employment. During your employment with the Company you will agree to work on any project that you are assigned to, irrespective of technical platforms / skills and nature of the project. If necessary, you may also be required to work in shifts.',False,False)],
        [('Regardless of any other M/s Cloud R tech entities or where you may be required to work overseas for any such M/s Cloud R tech entities for an extensive period, you shall at all times remain an employee of the Company exclusively and shall not be entitled to any such foreign salary or benefits (including medical insurance, green card sponsorship, etc.) payable or applicable to employees of such other M/s Cloud R tech entities other than the salary and benefits specified in this offer letter.',False,False)],
        [('You also agree that if you breach any of the terms and conditions stipulated in this Agreement, you will be liable for any loss or damage suffered directly or indirectly by the Company as a result of your action.',False,False)],
        [('You will be responsible for safekeeping and return in good condition and order of all company property, which may be in your use, custody or charge.',False,False)],
        [("The Company shall be at liberty to amend the whole or any part of this Agreement after its execution if it considers it necessary and reasonable but only after two (2) weeks' notice is duly given to you informing you of the proposed amendment(s). The terms and conditions of your proposed employment with the company in this letter supersede any contrary verbal representations concerning conditions of employment.",False,False)],
    ]
    for i,parts in enumerate(clauses,1):
        pc=doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        _para_spacing(pc,before=0,after=6,line=17.5); _run(pc,f'{i}.  ')
        for txt,b,it in parts: _run(pc,txt,bold=b,italic=it)
        _add_body_background_after_break(doc, LETTERHEAD, PAGE_W, PAGE_H, doc_pr_id=150+i)

    pc2=doc.add_paragraph(); pc2.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    _para_spacing(pc2,before=6,after=8,line=17.5)
    _run(pc2,'We would appreciate you, confirming your acceptance by signing on the space provided and returning this letter to HR department, indicating your proposed start date. Let me close by reaffirming our belief that the skill and background you bring to ')
    _run(pc2,'M/s Cloud R tech',bold=True)
    _run(pc2,' will be instrumental to the future success of the Company. Without hesitation, the single most important factor in our success has been our people. We look forward to working with you very soon.')

    pw=doc.add_paragraph(); _para_spacing(pw,before=0,after=12,line=17.5)
    _run(pw,'We welcome you to '); _run(pw,'M/s Cloud R tech',bold=True)
    _run(pw,' family and look forward to a fruitful collaboration. With best wishes,')

    # Safety checkpoint: the 12 clauses above reliably run past one page,
    # so the overflow page (natural pagination, no explicit break) would
    # otherwise have no background. Drop a zero-height background anchor
    # here so whichever page this paragraph lands on gets covered too.
    _add_body_background_after_break(doc, LETTERHEAD, PAGE_W, PAGE_H, doc_pr_id=150)

    _sig_block(doc,hr_name=hr_name,sig_src=sig_src)

    # Run-level page break (not doc.add_page_break()) — the latter adds an
    # extra empty paragraph that can push content onto an unwanted extra page.
    last_body_para = doc.paragraphs[-1]
    run_br = last_body_para.add_run()
    run_br.add_break(WD_BREAK.PAGE)

    _add_body_background_after_break(doc, LETTERHEAD, PAGE_W, PAGE_H, doc_pr_id=101)
    _annexure(doc,sal,name,designation,ctc_lpa,ctc_monthly)
    return _finalise_letter(doc, bg_path=LETTERHEAD)


# ═══════════════════════════════════════════════════════════════════════════════
# 4. CONTRACT EXTENSION
# ═══════════════════════════════════════════════════════════════════════════════
def generate_contract(data):
    doc=_setup_letter_doc()
    _add_body_background(doc, LETTERHEAD, PAGE_W, PAGE_H, doc_pr_id=100)
    ref=data['ref']; name=data['name']; sal=_salutation(data.get('gender','male'))
    address=data['address']; pincode=data['pincode']
    ext_date=data['extended_date']
    cur_lpa=data['current_ctc_lpa'].replace('LPA','').replace('lpa','').strip()
    inc_lpa=data['increment_ctc_lpa'].replace('LPA','').replace('lpa','').strip()
    eff_month=data['effective_month']; designation=data['designation']
    ctc_monthly=data['ctc_monthly']; hr_name=data.get('hr_name'); sig_src=data.get('sig_src')

    p=doc.add_paragraph(); _para_spacing(p,before=0,after=10,line=17.5)
    _run(p,f'Ref : {ref}',bold=True,italic=True)

    p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p2,before=0,after=8,line=17.5); _run(p2,'CONTRACT EXTENSION',bold=True,italic=True)

    p3=doc.add_paragraph(); _para_spacing(p3,before=0,after=0,line=17.5); _run(p3,'To:',bold=True,italic=True)
    _add_para(doc,f'{sal} {name},',before=0,after=0)
    _add_para(doc,address,before=0,after=0)
    _add_para(doc,f'{pincode}.',before=0,after=14)
    _add_para(doc,f'Dear {sal} {name},',before=0,after=12)

    p4=doc.add_paragraph(); p4.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    _para_spacing(p4,before=0,after=10,line=17.5)
    _run(p4,'We are glad to inform you that your Contract period is extended till ')
    _run(p4,ext_date,bold=True); _run(p4,' . Your remuneration has been revised from ')
    _run(p4,f'INR {cur_lpa} LPA to INR {inc_lpa} LPA',bold=True)
    _run(p4,f' with effect from the beginning of {eff_month}. A detailed of salary structure is given in the annexure')

    _add_para(doc,'All the other terms and conditions of the original contract remain unchanged.',before=0,after=24)
    _sig_block(doc,hr_name=hr_name,sig_src=sig_src)
    doc.add_page_break()
    _add_body_background_after_break(doc, LETTERHEAD, PAGE_W, PAGE_H, doc_pr_id=101)
    _annexure(doc,sal,name,designation,inc_lpa,ctc_monthly)
    return _finalise_letter(doc, bg_path=LETTERHEAD)


# ═══════════════════════════════════════════════════════════════════════════════
# 5. PAYSLIP — clean white document, NO background.
#    Table column widths match the PDF generator's exact pt measurements
#    (see pdf_generator.py generate_payslip) so DOCX and PDF render identically.
# ═══════════════════════════════════════════════════════════════════════════════

# Solid black borders, slightly heavier than letter tables, for a crisp
# print-style look matching the PDF (avoids the thin grey-ish look Word/
# LibreOffice can render at sz=4).
def _ps_cell_borders(cell, top=True, bottom=True, left=True, right=True, sz='6'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')): tcPr.remove(old)
    tcB = OxmlElement('w:tcBorders')
    for side, show in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        el = OxmlElement(f'w:{side}')
        if show:
            el.set(qn('w:val'),'single'); el.set(qn('w:sz'),sz)
            el.set(qn('w:space'),'0');   el.set(qn('w:color'),'000000')
        else:
            el.set(qn('w:val'),'none');  el.set(qn('w:sz'),'0')
            el.set(qn('w:color'),'auto')
        tcB.append(el)
    tcPr.append(tcB)


def generate_payslip(data):
    doc = Document()
    section = doc.sections[0]
    section.page_width    = Pt(PS_PAGE_W)
    section.page_height   = Pt(PS_PAGE_H)
    section.left_margin   = Pt(PS_LEFT_M)
    section.right_margin  = Pt(PS_RIGHT_M)
    section.top_margin    = Pt(PS_TOP_M)
    section.bottom_margin = Pt(PS_BOT_M)
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)
    compat = doc.settings.element.find(qn('w:compat'))
    if compat is not None:
        for cs in compat.findall(qn('w:compatSetting')):
            if cs.get(qn('w:name')) == 'compatibilityMode':
                cs.set(qn('w:val'), '15')

    gross = float(data.get('gross_salary', 0))
    basic = round(gross*0.40,2); hra   = round(gross*0.20,2); med  = round(gross*0.10,2)
    perf  = round(gross*0.15,2); conv  = round(gross*0.10,2); mgmt = round(gross*0.05,2)
    total_earn = round(basic+hra+med+perf+conv+mgmt, 2)
    other_ded  = round(basic/12, 2); pt_tax = 250.00
    total_ded  = round(other_ded+pt_tax, 2); net_pay = round(total_earn-total_ded, 2)
    def rs(v): return f"Rs.{v:,.2f}"

    wdays      = str(data.get('working_days', ''))
    phol       = str(data.get('paid_holiday', '0'))
    month_year = data.get('month_year', '')

    L=WD_ALIGN_PARAGRAPH.LEFT; R=WD_ALIGN_PARAGRAPH.RIGHT; C=WD_ALIGN_PARAGRAPH.CENTER

    # ── Header bar image + Title bar — single continuous bordered block ────
    # Both live in the SAME 2-row table so there is zero gap and the outer
    # border is continuous: row-0 = banner image, row-1 = "Pay Slip for …".
    _tw = 532.82
    bt = doc.add_table(rows=2, cols=1)
    bt.alignment = WD_TABLE_ALIGNMENT.LEFT
    _fixed_table_layout(bt, _tw, [_tw])
    # Remove any cell spacing at table level so rows sit flush
    _tblPr = bt._tbl.tblPr
    for old in _tblPr.findall(qn('w:tblCellSpacing')):
        _tblPr.remove(old)

    # --- Row 0: banner image inside a bordered cell ---
    bc = bt.rows[0].cells[0]
    _col_width(bc, _tw)
    _ps_cell_borders(bc, top=True, bottom=True, left=True, right=True, sz='6')
    _cell_margins(bc, 0, 0, 0, 0)
    bp = bc.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(bp, before=0, after=0, line=0)
    if PS_HEADER:
        ps_buf, _ = _to_icc_free_jpeg(PS_HEADER)
        _img_w = _tw
        try:
            with PILImage.open(PS_HEADER) as _im:
                _iw, _ih = _im.size
            _img_h = _img_w * (_ih / _iw)
        except Exception:
            _img_h = 62.2
        run = bp.add_run()
        ps_buf.seek(0)
        run.add_picture(ps_buf, width=Pt(_img_w), height=Pt(_img_h))
        # Set row height to exact image height — eliminates any white gap
        _trPr = bt.rows[0]._tr.get_or_add_trPr()
        _trHeight = OxmlElement('w:trHeight')
        _trHeight.set(qn('w:val'), str(int(_img_h * 20)))
        _trHeight.set(qn('w:hRule'), 'exact')
        for old in _trPr.findall(qn('w:trHeight')):
            _trPr.remove(old)
        _trPr.append(_trHeight)

    # --- Row 1: title text inside a bordered cell ---
    tc = bt.rows[1].cells[0]
    _col_width(tc, _tw)
    _ps_cell_borders(tc, top=True, bottom=True, left=True, right=True, sz='6')
    _cell_margins(tc, 60, 60, 80, 80)
    tp = tc.paragraphs[0]
    tp.clear()
    tp.alignment = C
    _para_spacing(tp, before=0, after=0, line=15)
    _run(tp, f'Pay Slip for the month of {month_year}', bold=True, size=12)

    # ── Employee info block — column widths match PDF exactly ──────────────────
    # PDF group-local verticals: x=71.42, x=238.00, x=454.42, x=532.82
    ecw = [71.42, 166.58, 216.42, 78.40]
    et = doc.add_table(rows=2, cols=4); et.alignment = WD_TABLE_ALIGNMENT.LEFT
    _fixed_table_layout(et, sum(ecw), ecw)
    r0 = et.rows[0]
    for j,(lbl,val) in enumerate([
        ('Employee No', data.get('emp_no','')),
        ('Name',        data.get('name','')),
        ('Designation', data.get('designation','')),
        ('Category',    data.get('category',''))
    ]):
        c = r0.cells[j]; _col_width(c,ecw[j])
        _ps_cell_borders(c, top=True, bottom=True, left=True, right=True)
        _cell_margins(c,40,40,80,40)
        p = c.paragraphs[0]; p.clear(); _para_spacing(p,before=0,after=4,line=12)
        _run(p, lbl, bold=True, size=9)
        p.add_run().add_break()
        _run(p, val, size=10)

    # Row 2: [EmpNo+Name merged, blank] | Sex | Date of Joining (matches PDF: no divider in DOJ box)
    r1 = et.rows[1]
    mg = r1.cells[0].merge(r1.cells[1]); _col_width(mg, ecw[0]+ecw[1])
    _ps_cell_borders(mg, top=False, bottom=True, left=True, right=True)
    _cell_margins(mg,40,40,80,40)
    mg.paragraphs[0].clear(); _para_spacing(mg.paragraphs[0],before=0,after=0,line=12)

    sx = r1.cells[2]; _col_width(sx, 122.0)
    _ps_cell_borders(sx, top=False, bottom=True, left=True, right=True)
    _cell_margins(sx,40,40,80,40)
    sp = sx.paragraphs[0]; sp.clear(); _para_spacing(sp,before=0,after=0,line=12)
    _run(sp,'Sex: ',bold=True,size=9); _run(sp,data.get('sex',''),size=10)

    dj = r1.cells[3]; _col_width(dj, ecw[2]+ecw[3]-122.0)
    _ps_cell_borders(dj, top=False, bottom=True, left=False, right=True)
    _cell_margins(dj,40,40,60,30)
    dp = dj.paragraphs[0]; dp.clear(); _para_spacing(dp,before=0,after=0,line=12)
    _run(dp,'Date of Joining: ',bold=True,size=9); _run(dp,data.get('date_of_joining',''),size=10)

    # ── Salary table — column widths match PDF exactly ─────────────────────────
    # PDF group-local verticals: x=103.70, x=237.53, x=320.59, x=450.39, x=532.82
    #   Attendance-label 71.42 | Attendance-value 32.28
    #   Earnings particulars 133.83 | Earnings amount 83.06
    #   Deductions particulars 129.80 | Deductions amount 82.43
    scw = [71.42, 32.28, 133.83, 83.06, 129.80, 82.43]
    sal = doc.add_table(rows=10, cols=6); sal.alignment = WD_TABLE_ALIGNMENT.LEFT
    _fixed_table_layout(sal, sum(scw), scw)

    def sc(ri,ci,txt,bold=False,align=L,sz=10,
           top=True,bottom=True,left=True,right=True):
        cell=sal.rows[ri].cells[ci]; _col_width(cell,scw[ci])
        _ps_cell_borders(cell, top=top, bottom=bottom, left=left, right=right)
        _cell_margins(cell,16,16,40,40); p=cell.paragraphs[0]; p.clear()
        p.alignment=align; _para_spacing(p,before=0,after=0,line=13)
        _run(p,txt,bold=bold,size=sz)

    # Header row: "Working Days" | (blank) | "Earnings" (spans particulars+amount) | "Deductions" (spans)
    sc(0,0,'Working Days',bold=True)
    sc(0,1,'')
    sc(0,2,'Earnings',bold=True,align=C)
    sc(0,3,'')
    # Merge cells 4+5 for "Deductions" header so no vertical line inside the cell
    dm = sal.rows[0].cells[4].merge(sal.rows[0].cells[5])
    _ps_cell_borders(dm); _cell_margins(dm,16,16,40,40)
    dmp = dm.paragraphs[0]; dmp.clear(); dmp.alignment = C
    _para_spacing(dmp,before=0,after=0,line=13); _run(dmp,'Deductions',bold=True,size=10)

    sc(1,0,'Attendance',bold=True)
    sc(1,1,'')
    sc(1,2,'Particulars',bold=True)
    sc(1,3,'Amount',bold=True,align=R)
    sc(1,4,'Particulars',bold=True)
    sc(1,5,'Amount',bold=True,align=R)

    srows=[
        ('Working Days', wdays, 'Basic',               rs(basic), 'Other Deduction',  rs(other_ded)),
        ('Paid Holiday', phol,  'HRA',                 rs(hra),   'Professional Tax', rs(pt_tax)),
        ('','',                 'Medical Allowance',    rs(med),   '',''),
        ('','',                 'Performance Bonus',    rs(perf),  '',''),
        ('','',                 'Conveyance',           rs(conv),  '',''),
        ('','',                 'Management Allowance', rs(mgmt),  '',''),
    ]
    for i,(a,b,c,d,e,f) in enumerate(srows):
        ri = i+2
        sc(ri,0,a); sc(ri,1,b,align=R)
        sc(ri,2,c); sc(ri,3,d,align=R); sc(ri,4,e); sc(ri,5,f,align=R)

    sc(8,0,'Total Days',bold=True); sc(8,1,wdays,bold=True,align=R)
    sc(8,2,'Total Earnings',bold=True); sc(8,3,rs(total_earn),bold=True,align=R)
    sc(8,4,'Total Deductions',bold=True); sc(8,5,rs(total_ded),bold=True,align=R)

    # Net Pay row — label spans first 5 columns, amount in last column (matches PDF)
    nr = sal.rows[9]; nm = nr.cells[0]
    for k in range(1,5): nm = nm.merge(nr.cells[k])
    _ps_cell_borders(nm); _cell_margins(nm,16,16,40,40)
    np_ = nm.paragraphs[0]; np_.clear(); np_.alignment = R
    _para_spacing(np_,before=0,after=0,line=13); _run(np_,'Net Pay:',bold=True,size=10)
    na = nr.cells[5]; _col_width(na,scw[5]); _ps_cell_borders(na); _cell_margins(na,16,16,40,40)
    nap = na.paragraphs[0]; nap.clear(); nap.alignment = R
    _para_spacing(nap,before=0,after=0,line=13); _run(nap,rs(net_pay),bold=True,size=10)

    # ── Note ─────────────────────────────────────────────────────────────────
    note = doc.add_paragraph(); _para_spacing(note,before=8,after=0,line=12)
    _run(note,'NOTE: This is a computer generated copy and it does not need any seal or signature.',bold=True,size=9)

    # No background — clean white payslip. Colour post-process still applied
    # so the header banner's orange prints solid instead of Word dulling it.
    buf = BytesIO()
    doc.save(buf)
    return _postprocess_docx_color(buf.getvalue())